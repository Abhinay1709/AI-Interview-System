import re
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt

# ==========================================================
# COLOUR PALETTE
# ==========================================================

COLOR_DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # headings
COLOR_MID_BLUE    = RGBColor(0x2E, 0x74, 0xB5)    # sub-headings / stats
COLOR_ORANGE      = RGBColor(0xC0, 0x50, 0x00)    # feedback label
COLOR_DARK_GRAY   = RGBColor(0x40, 0x40, 0x40)    # body labels
COLOR_MID_GRAY    = RGBColor(0x70, 0x70, 0x70)    # secondary text
COLOR_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GREEN       = RGBColor(0x37, 0x5A, 0x2E)    # strengths heading

HEX_HEADER_DARK   = "1F497D"    # dark-blue table header bg
HEX_HEADER_MID    = "2E74B5"    # mid-blue table header bg
HEX_CELL_LIGHT    = "DEEAF1"    # light blue cell bg (scores)
HEX_CELL_MID      = "EBF3FB"    # very light blue cell bg (stats)
HEX_SEPARATOR     = "BBBBBB"    # horizontal-rule colour


# ==========================================================
# HELPERS — XML / STYLING
# ==========================================================

def _set_cell_bg(cell, hex_color):
    """Apply a solid background fill to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_padding(cell,
                    top=80, bottom=80,
                    left=120, right=120):
    """Set inner cell margins (EMU-ish units in docx = twentieths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (
        ("top",    top),
        ("bottom", bottom),
        ("left",   left),
        ("right",  right),
    ):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_horizontal_rule(doc, color=HEX_SEPARATOR):
    """Draw a thin 1-pt bottom-border paragraph as a visual divider."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(6)
    pPr   = para._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    btm   = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "4")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), color)
    pBdr.append(btm)
    pPr.append(pBdr)


def _spacer(doc, before=4, after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    return para


# ==========================================================
# HELPERS — CONTENT BLOCKS
# ==========================================================

def _add_section_heading(doc, text, level=1):
    """
    Bold section heading.
      level 1 — 14pt dark-blue  (top-level section)
      level 2 — 12pt mid-blue   (question header)
      level 3 — 11pt dark-gray  (sub-label)
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)

    run = para.add_run(text)
    run.bold = True

    if level == 1:
        run.font.size  = Pt(14)
        run.font.color.rgb = COLOR_DARK_BLUE
    elif level == 2:
        run.font.size  = Pt(12)
        run.font.color.rgb = COLOR_MID_BLUE
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = COLOR_DARK_GRAY
    return para


def _add_label_value(doc, label, value, label_color=None):
    """One line with a **bold label** then plain value."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)

    lr = para.add_run(f"{label}:  ")
    lr.bold = True
    lr.font.size = Pt(11)
    lr.font.color.rgb = label_color or COLOR_DARK_GRAY

    vr = para.add_run(str(value) if value else "—")
    vr.font.size = Pt(11)
    return para


def _add_block_content(doc, label, text, label_color=None, indent=True):
    """Bold label on its own line, then indented body text."""
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after  = Pt(1)
    lr = lp.add_run(f"{label}:")
    lr.bold = True
    lr.font.size  = Pt(11)
    lr.font.color.rgb = label_color or COLOR_DARK_GRAY

    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(1)
    cp.paragraph_format.space_after  = Pt(4)
    if indent:
        cp.paragraph_format.left_indent = Inches(0.2)
    cr = cp.add_run(str(text) if text else "—")
    cr.font.size = Pt(11)
    return cp


def _add_bullet_points(doc, text):
    
    if not text:
        return

    text = str(text)

    text = text.replace(
        "• ",
        "\n• "
    )

    text = re.sub(
        r"\n+",
        "\n",
        text
    )

    for line in text.splitlines():

        line = re.sub(
            r"^[•\-\*]\s*",
            "",
            line
        ).strip()

        if not line:
            continue
        
        para=doc.add_paragraph()
        #para.paragraph_format.left_indent = Inches(0.25)
        #para.paragraph_format.space_after = Pt(2)
        
        run=para.add_run(f"• {line}")
        
        run.font.size = Pt(11)

# ==========================================================
# HELPERS — TWO-ROW SUMMARY TABLE
# ==========================================================

def _build_summary_table(doc, headers, values,
                        hdr_hex, cell_hex,
                        hdr_color=None, val_color=None):
    """
    Create a 2-row (header + value) table.
    headers / values must have the same length (2 or 4 cols).
    """
    n   = len(headers)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.style = "Table Grid"

    # Centre the whole table
    tblEl = tbl._tbl
    tPr   = tblEl.find(qn("w:tblPr"))
    if tPr is None:
        tPr = OxmlElement("w:tblPr")
        tblEl.insert(0, tPr)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tPr.append(jc)

    for i, (hdr, val) in enumerate(zip(headers, values)):

        # ---- header row ----
        hc = tbl.rows[0].cells[i]
        _set_cell_bg(hc, hdr_hex)
        _set_cell_padding(hc)
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(hdr)
        hr.bold = True
        hr.font.size  = Pt(10)
        hr.font.color.rgb = hdr_color or COLOR_WHITE

        # ---- value row ----
        vc = tbl.rows[1].cells[i]
        _set_cell_bg(vc, cell_hex)
        _set_cell_padding(vc, top=120, bottom=120)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(str(val))
        vr.bold = True
        vr.font.size  = Pt(15)
        vr.font.color.rgb = val_color or COLOR_MID_BLUE

    return tbl


# ==========================================================
# REGEX HELPERS (kept identical to original)
# ==========================================================

def safe_extract(pattern, text, default="Not Available"):
    try:
        match = re.search(
            pattern,
            str(text),
            re.IGNORECASE | re.DOTALL
        )
        if match:
            value = match.group(1).strip()
            if value:
                return value
    except Exception:
        pass
    return default


def extract_question_score(evaluation_text, question_number):
    return safe_extract(
        rf"Question\s+{question_number}\s+Score:\s*(\d+\/10)",
        evaluation_text,
        "0/10"
    )


def extract_model_answer(evaluation_text, question_number):
    try:
        pattern = (
            rf"Question\s+{question_number}\s+Score:.*?"
            rf"Model Answer:\s*(.*?)"
            rf"Feedback:"
        )
        match = re.search(
            pattern, evaluation_text,
            re.IGNORECASE | re.DOTALL
        )
        if match:
            return match.group(1).strip()
    except Exception:
        pass
    return "Not Available"


def extract_feedback(evaluation_text, question_number):
    try:
        pattern = (
            rf"Question\s+{question_number}\s+Score:.*?"
            rf"Feedback:\s*(.*?)"
            rf"(Question\s+\d+\s+Score:|Technical Score:)"
        )
        match = re.search(
            pattern, evaluation_text,
            re.IGNORECASE | re.DOTALL
        )
        if match:
            return match.group(1).strip()
    except Exception:
        pass
    return "Not Available"


def generate_score_chart(
    technical,
    communication,
    confidence,
    overall
):

    plt.figure(
        figsize=(5,3)
    )

    labels = [

        "Technical",
        "Communication",
        "Confidence",
        "Overall"

    ]

    values = [

        technical,
        communication,
        confidence,
        overall

    ]

    plt.bar(
        labels,
        values
    )

    plt.ylim(
        0,
        10
    )

    chart_path = (
        "score_chart.png"
    )

    plt.savefig(
        chart_path,
        bbox_inches="tight"
    )

    plt.close()

    return chart_path
# ==========================================================
# MAIN REPORT GENERATOR
# ==========================================================

def generate_full_report(questions, answers, evaluation):
    """
    Build a professional .docx report and return it as bytes.
    Compatible with Streamlit's st.download_button(data=...).
    """
    # ---- sanitise evaluation text ----
    try:
        from modules.answer_evaluator import clean_evaluation_text
        evaluation = clean_evaluation_text(evaluation)
    except Exception:
        pass

    doc = Document()

    # ----------------------------------------------------------
    # PAGE SETUP — US Letter, 1-inch margins
    # ----------------------------------------------------------
    sec = doc.sections[0]
    sec.page_width      = Inches(8.5)
    sec.page_height     = Inches(11)
    sec.top_margin      = Inches(1.0)
    sec.bottom_margin   = Inches(1.0)
    sec.left_margin     = Inches(1.0)
    sec.right_margin    = Inches(1.0)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ==========================================
    # COVER PAGE
    # ==========================================
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(
        "AI INTERVIEW COACH"
    )
    run.bold = True
    run.font.size = Pt(26)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Professional Interview Assessment Report"
    )
    subtitle_run.font.size = Pt(14)
    doc.add_paragraph("")
    # Candidate Information
    candidate_name = "Candidate"
    info_table = doc.add_table(
        rows=5,
        cols=2
    )
    info_table.style = "Table Grid"
    info_table.cell(0,0).text = "Candidate Name"
    info_table.cell(0,1).text = candidate_name
    info_table.cell(1,0).text = "Interview Date"
    info_table.cell(1,1).text = datetime.now().strftime(
        "%d %B %Y"
    )
    overall_score = safe_extract(
        r"Overall Score:\s*([0-9]+(?:\/10)?)",
        evaluation,
        "N/A"
    )
    info_table.cell(2,0).text = "Overall Score"
    info_table.cell(2,1).text = overall_score
    technical_score = safe_extract(
        r"Technical Score:\s*([0-9]+(?:\/10)?)",
        evaluation,
        "N/A"
    )
    info_table.cell(3,0).text = "Technical Score"
    info_table.cell(3,1).text = technical_score
    readiness = "Placement Ready"
    info_table.cell(4,0).text = "Assessment"
    info_table.cell(4,1).text = readiness
    doc.add_paragraph("")
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(
        "This report contains interview performance analysis, "
        "question-wise evaluation, strengths, weaknesses, "
        "and recommendations for improvement."
    )
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        "Generated by AI Interview Coach"
    )
    _add_horizontal_rule(
        doc,
        color="1F497D"
    )    


    # ==========================================
    # EXECUTIVE SUMMARY
    # ==========================================

    _add_section_heading(
        doc,
        "EXECUTIVE SUMMARY"
    )

    overall_score = safe_extract(
        r"Overall Score:\s*(\d+)",
        evaluation,
        "0"
    )

    try:

        score = int(
            overall_score
        )

    except:

        score = 0

    if score >= 8:
        summary_text = (
            "The candidate demonstrated strong technical knowledge, "
            "good communication skills, and confidence throughout the interview. "
            "Overall performance indicates readiness for most technical interview opportunities."
        )
    elif score >= 6:
        summary_text = (
            "The candidate showed a good understanding of core concepts but has "
            "areas that require further improvement. With additional practice and "
            "preparation, interview performance can be significantly enhanced."
        )
    else:
        summary_text = (
            "The candidate requires further preparation in technical concepts, "
            "communication, and interview confidence before attending professional interviews."
        )

    doc.add_paragraph(
        summary_text
    )

    _add_horizontal_rule(
        doc
    )

    # ==========================================================
    # SECTION 1 — OVERALL EVALUATION SCORES
    # ==========================================================

    _add_section_heading(doc, "OVERALL EVALUATION SCORES")

    technical_score    = safe_extract(r"Technical Score:\s*(\d+\/10)",    evaluation, "N/A")
    communication_score= safe_extract(r"Communication Score:\s*(\d+\/10)", evaluation, "N/A")
    confidence_score   = safe_extract(r"Confidence Score:\s*(\d+\/10)",   evaluation, "N/A")
    overall_score      = safe_extract(r"Overall Score:\s*(\d+\/10)",      evaluation, "N/A")

    _build_summary_table(
        doc,
        headers=["Technical Score", "Communication Score",
                "Confidence Score", "Overall Score"],
        values=[technical_score, communication_score,
                confidence_score, overall_score],
        hdr_hex=HEX_HEADER_DARK,
        cell_hex=HEX_CELL_LIGHT,
    )

    _add_horizontal_rule(
        doc,
        color="1F497D"
    )

    # --- BUG FIX APPLIED HERE ---
    def parse_chart_score(score_str):
        """Safely extracts the numerator for the chart, defaulting to 0.0 if 'N/A'."""
        try:
            return float(score_str.split("/")[0])
        except (ValueError, AttributeError):
            return 0.0

    chart = generate_score_chart(
        parse_chart_score(technical_score),
        parse_chart_score(communication_score),
        parse_chart_score(confidence_score),
        parse_chart_score(overall_score)
    )

    doc.add_picture(
        chart,
        width=Inches(4.5)
    )
    
    _add_horizontal_rule(
        doc,
        color="1F497D"
    )
    
    _spacer(doc, before=10, after=4)
    
    # ==========================================================
    # SECTION 2 — INTERVIEW STATISTICS
    # ==========================================================

    _add_section_heading(doc, "INTERVIEW STATISTICS")

    attempted_match = re.search(
        r"Questions Attempted:\s*(\d+)\/(\d+)",
        evaluation, re.IGNORECASE
    )
    skipped_match = re.search(
        r"Questions Skipped:\s*(\d+)\/(\d+)",
        evaluation, re.IGNORECASE
    )

    total_q   = attempted_match.group(2) if attempted_match else str(len(questions))
    attempted = attempted_match.group(1) if attempted_match else "0"
    skipped   = skipped_match.group(1)   if skipped_match   else "0"

    completion_pct = "0%"
    try:
        if attempted_match:
            a = int(attempted_match.group(1))
            t = int(attempted_match.group(2))
            if t > 0:
                completion_pct = f"{round((a / t) * 100, 1)}%"
    except Exception:
        pass

    _build_summary_table(
        doc,
        headers=["Total Questions", "Answered", "Skipped", "Completion %"],
        values=[total_q, attempted, skipped, completion_pct],
        hdr_hex=HEX_HEADER_MID,
        cell_hex=HEX_CELL_MID,
    )
    
    _add_horizontal_rule(
        doc,
        color="1F497D"
    )

    _spacer(doc, before=8, after=4)

    _add_horizontal_rule(
        doc,
        color="1F497D"
    )

    # ==========================================================
    # SECTION 3 — QUESTION-WISE ANALYSIS
    # ==========================================================

    _add_section_heading(doc, "QUESTION-WISE ANALYSIS")
    _add_horizontal_rule(
        doc,
        color="000000"
    )

    if not questions:
        doc.add_paragraph("No questions available.")
    else:
        for idx, question in enumerate(questions, start=1):

            answer       = answers.get(question, "No Answer")
            score        = extract_question_score(evaluation, idx)
            model_answer = extract_model_answer(evaluation, idx)
            feedback     = extract_feedback(evaluation, idx)

            # ---- Question header (bold, dark-blue) ----
            qh = doc.add_paragraph()
            qh.paragraph_format.space_before = Pt(12)
            qh.paragraph_format.space_after  = Pt(3)
            qhr = qh.add_run(f"Question {idx}   |   Score: {score}")
            qhr.bold       = True
            qhr.font.size  = Pt(12)
            qhr.font.color.rgb = COLOR_DARK_BLUE

            # ---- Question text ----
            _add_label_value(doc, "Question", question)

            # ---- My Answer ----
            _add_label_value(doc, "My Answer", answer)

            # ---- Model Answer ----
            _add_block_content(
                doc, "Model Answer", model_answer,
                label_color=COLOR_MID_BLUE
            )

            # ---- Feedback ----
            _add_block_content(
                doc, "Feedback", feedback,
                label_color=COLOR_ORANGE
            )

            # ---- Thin separator between questions ----
            _add_horizontal_rule(doc)

    _add_horizontal_rule(
        doc,
        color="1F497D"
    )
    # ==========================================================
    # SECTION 4 — STRENGTHS
    # ==========================================================

    _add_section_heading(doc, "STRENGTHS")

    strengths = safe_extract(
        r"Strengths:(.*?)Weaknesses:",
        evaluation, "Not Available"
    )
    _add_bullet_points(doc, strengths)
    
    _add_horizontal_rule(
        doc,
        color="1F497D"
    )

    _spacer(doc, before=6, after=4)

    # ==========================================================
    # SECTION 5 — WEAKNESSES
    # ==========================================================

    _add_section_heading(doc, "WEAKNESSES")

    weaknesses = safe_extract(
        r"Weaknesses:(.*?)Suggestions:",
        evaluation, "Not Available"
    )
    _add_bullet_points(doc, weaknesses)

    _add_horizontal_rule(
        doc,
        color="1F497D"
    )
    _spacer(doc, before=6, after=4)

    # ==========================================
    # READINESS ASSESSMENT
    # ==========================================

    _add_section_heading(
        doc,
        "READINESS ASSESSMENT"
    )

    try:

        overall_num = int(
            overall_score.split("/")[0]
        )

    except:

        overall_num = 0

    readiness = overall_num * 10

    _add_label_value(
        doc,
        "Readiness Score",
        f"{readiness}%"
    )

    if readiness >= 80:

        doc.add_paragraph(
            "Placement Ready"
        )

    elif readiness >= 60:

        doc.add_paragraph(
            "Almost Ready"
        )

    else:

        doc.add_paragraph(
            "Needs More Preparation"
        )

    _add_horizontal_rule(
        doc
    )

    # ==========================================
    # FINAL RECOMMENDATION
    # ==========================================
    _add_section_heading(
        doc,
        "FINAL RECOMMENDATION"
    )
    if readiness >= 80:
        recommendation = (
            "Candidate is ready for interview opportunities."
        )
    elif readiness >= 60:
        recommendation = (
            "Candidate should improve weak areas before attending interviews."
        )
    else:
        recommendation = (
            "Candidate should focus on fundamentals and practice interviews."
        )
    doc.add_paragraph(
        recommendation
    )
    _add_horizontal_rule(
        doc
    )

    # ==========================================================
    # SECTION 6 — SUGGESTIONS
    # ==========================================================

    _add_section_heading(doc, "SUGGESTIONS FOR IMPROVEMENT")

    suggestions = safe_extract(
        r"Suggestions:(.*)",
        evaluation, "Not Available"
    )
    _add_bullet_points(doc, suggestions)

    _spacer(doc, before=12, after=4)

    # ==========================================================
    # FOOTER LINE
    # ==========================================================

    _add_horizontal_rule(doc, color="2E74B5")

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(2)
    fr = fp.add_run(
        "AI Interview Preparation System  |  Abhinay Andhavarapu"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = COLOR_MID_GRAY
    fr.italic = True

    # ==========================================================
    # SERIALISE TO BYTES
    # ==========================================================

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()